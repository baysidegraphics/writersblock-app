# WB_BUILD_PROJECT_TITLE_SINGLE_SAVED_DROPDOWN_0620
import streamlit as st
import re
import time
import os
import json
from dotenv import load_dotenv
from datetime import datetime, timezone
from io import BytesIO
from docx import Document
from docx.shared import Pt
from gemini_backend import GeminiWriterBackend
from supabase import create_client
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
def get_current_user_id():
    result = supabase.table("writersblock_users").select("id").limit(1).execute()

    if result.data and len(result.data) > 0:
        return result.data[0]["id"]

    return None


def save_project_to_supabase(project_title, project_type="story"):
    user_id = get_current_user_id()

    if not user_id:
        st.error("No user found.")
        return None

    result = supabase.table("writersblock_projects").insert({
        "user_id": user_id,
        "project_title": project_title,
        "project_type": project_type,
        "project_status": "active"
    }).execute()

    if result.data and len(result.data) > 0:
        st.session_state.current_project_id = result.data[0].get("id")

    return result

    def save_draft_to_supabase(project_id, draft_title, draft_content, draft_type="sandbox"):
        user_id = get_current_user_id()

    if not user_id:
        st.error("No user found.")
        return None

    result = supabase.table("writersblock_project_drafts").insert({
        "project_id": project_id,
        "user_id": user_id,
        "draft_title": draft_title,
        "draft_type": draft_type,
        "draft_content": draft_content,
        "version_number": 1
    }).execute()

def save_project_draft_to_supabase(project_id, draft_title, draft_type, draft_content):
    user_id = get_current_user_id()

    if not user_id:
        st.error("No user found.")
        return None

    result = supabase.table("writersblock_project_drafts").insert({
        "project_id": project_id,
        "user_id": user_id,
        "draft_title": draft_title,
        "draft_type": draft_type,
        "draft_content": draft_content,
        "version_number": 1
    }).execute()

    return result


def get_project_session_snapshot():
    snapshot_keys = [
        "project_stage",
        "story_seed",
        "project_title",
        "format_type",
        "tone_style",
        "author_style",
        "sandbox_slates",
        "chosen_option",
        "chosen_option_text",
        "plot_outline",
        "current_draft_output",
        "rolling_lore_book",
        "completed_beats",
        "scene_blocks",
        "current_processing_beat",
        "manuscript_version",
        "rewritten_manuscript",
        "rewrite_applied_message",
        "rewrite_preview_id",
        "rewrite_target_length",
        "writer_credits_remaining",
        "plot_outline_credit_pending",
        "plot_outline_credit_cost",
        "charged_scene_blocks",
        "revision_report",
        "export_view",
        "current_project_id",
    ]

    snapshot = {}
    for key in snapshot_keys:
        value = st.session_state.get(key)
        if key == "scene_blocks" and isinstance(value, dict):
            snapshot[key] = {str(k): v for k, v in value.items()}
        else:
            snapshot[key] = value

    snapshot["saved_at"] = datetime.now(timezone.utc).isoformat()
    return snapshot


def save_current_project_session_to_supabase():
    user_id = get_current_user_id()

    if not user_id:
        st.error("No user found.")
        return None

    project_id = get_active_project_id()
    if not project_id:
        result = save_project_to_supabase(
            project_title=get_project_title_for_display()[:80],
            project_type=st.session_state.get("format_type", "story"),
        )
        project_id = st.session_state.get("current_project_id")

    if not project_id:
        st.error("No active project available to save.")
        return None

    ensure_project_title_loaded()

    snapshot = get_project_session_snapshot()
    snapshot["current_project_id"] = project_id
    st.session_state.last_saved_project_at = snapshot.get("saved_at")

    result = supabase.table("writersblock_project_drafts").insert({
        "project_id": project_id,
        "user_id": user_id,
        "draft_title": "Saved Project Session",
        "draft_type": "project_session",
        "draft_content": json.dumps(snapshot),
        "version_number": int(st.session_state.get("current_processing_beat", 1) or 1),
    }).execute()

    return result


def get_latest_saved_project_session():
    user_id = get_current_user_id()

    query = (
        supabase.table("writersblock_project_drafts")
        .select("id, project_id, draft_content, created_at")
        .eq("draft_type", "project_session")
        .order("created_at", desc=True)
        .limit(1)
    )

    if user_id:
        query = query.eq("user_id", user_id)

    result = query.execute()

    if result.data and len(result.data) > 0:
        return result.data[0]

    return None


def load_saved_project_session_to_state(saved_session_row):
    if not saved_session_row:
        st.warning("No saved project session found.")
        return False

    try:
        snapshot = json.loads(saved_session_row.get("draft_content") or "{}")
    except Exception as e:
        st.error(f"Saved project could not be loaded: {e}")
        return False

    for key, value in snapshot.items():
        if key == "saved_at":
            st.session_state.last_saved_project_at = value
            continue
        if key == "scene_blocks" and isinstance(value, dict):
            restored = {}
            for scene_key, scene_text in value.items():
                try:
                    restored[int(scene_key)] = scene_text
                except Exception:
                    restored[scene_key] = scene_text
            st.session_state.scene_blocks = restored
        else:
            st.session_state[key] = value

    st.session_state.current_project_id = saved_session_row.get("project_id") or snapshot.get("current_project_id")
    return True


def save_manuscript_version_to_supabase(manuscript_content, version_label="Accepted Rewrite", source_stage="rewrite_engine"):
    user_id = get_current_user_id()

    if not user_id:
        st.error("No user found.")
        return None

    project_id = get_active_project_id()

    if not project_id:
        st.error("No active project ID found. Start a new project before saving manuscript versions.")
        return None

    existing_versions = (
        supabase.table("writersblock_manuscript_versions")
        .select("version_number")
        .eq("project_id", project_id)
        .order("version_number", desc=True)
        .limit(1)
        .execute()
    )

    if existing_versions.data:
        next_version = int(existing_versions.data[0]["version_number"]) + 1
    else:
        next_version = 1

    result = (
        supabase.table("writersblock_manuscript_versions")
        .insert({
            "project_id": project_id,
            "user_id": user_id,
            "version_number": next_version,
            "version_label": f"{get_project_title_for_display()} — {version_label}",
            "manuscript_content": manuscript_content,
            "source_stage": source_stage
        })
        .execute()
    )

    return result
    user_id = get_current_user_id()

    if not user_id:
        st.error("No user found.")
        return None

    latest_project = (
        supabase.table("writersblock_projects")
        .select("id")
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    )

    if not latest_project.data:
        st.error("No project found.")
        return None

    project_id = latest_project.data[0]["id"]

    existing_versions = (
        supabase.table("writersblock_manuscript_versions")
        .select("version_number")
        .eq("project_id", project_id)
        .order("version_number", desc=True)
        .limit(1)
        .execute()
    )

    if existing_versions.data:
        next_version = int(existing_versions.data[0]["version_number"]) + 1
    else:
        next_version = 1

    result = (
        supabase.table("writersblock_manuscript_versions")
        .insert({
            "project_id": project_id,
            "user_id": user_id,
            "version_number": next_version,
            "version_label": version_label,
            "manuscript_content": manuscript_content,
            "source_stage": source_stage
        })
        .execute()
    )

    return result



def get_active_project_id():
    project_id = st.session_state.get("current_project_id")
    if project_id:
        return project_id

    latest_project = (
        supabase.table("writersblock_projects")
        .select("id")
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    )

    if latest_project.data and len(latest_project.data) > 0:
        project_id = latest_project.data[0].get("id")
        st.session_state.current_project_id = project_id
        return project_id

    return None


def get_project_title_from_project_id(project_id):
    if not project_id:
        return ""

    try:
        result = (
            supabase.table("writersblock_projects")
            .select("project_title")
            .eq("id", project_id)
            .limit(1)
            .execute()
        )

        if result.data and len(result.data) > 0:
            return str(result.data[0].get("project_title") or "").strip()
    except Exception:
        return ""

    return ""


def ensure_project_title_loaded():
    if get_project_title():
        return True

    project_id = st.session_state.get("current_project_id")
    title_from_project = get_project_title_from_project_id(project_id)

    if title_from_project:
        st.session_state.project_title = title_from_project
        return True

    return False


def get_manuscript_versions_for_current_project():
    result = (
        supabase.table("writersblock_manuscript_versions")
        .select("id, version_number, version_label, manuscript_content, created_at, project_id, user_id, writersblock_projects(project_title)")
        .order("created_at", desc=False)
        .execute()
    )

    return result.data or []


def delete_manuscript_version_ids(version_ids):
    deleted_count = 0

    for version_id in version_ids:
        if not version_id:
            continue

        supabase.table("writersblock_manuscript_versions").delete().eq("id", version_id).execute()
        deleted_count += 1

    return deleted_count


def remove_duplicate_manuscript_versions_for_current_project():
    versions = get_manuscript_versions_for_current_project()
    seen = {}
    delete_ids = []

    sorted_versions = sorted(
        versions,
        key=lambda row: str(row.get("created_at") or ""),
        reverse=True,
    )

    for version in sorted_versions:
        content_key = re.sub(r"\s+", " ", str(version.get("manuscript_content") or "").strip())

        if not content_key:
            delete_ids.append(version.get("id"))
            continue

        if content_key in seen:
            delete_ids.append(version.get("id"))
        else:
            seen[content_key] = version.get("id")

    return delete_manuscript_version_ids(delete_ids)


def purge_old_manuscript_versions_for_current_project(keep_latest_count=3):
    versions = get_manuscript_versions_for_current_project()

    sorted_versions = sorted(
        versions,
        key=lambda row: str(row.get("created_at") or ""),
        reverse=True,
    )

    keep_latest_count = max(1, int(keep_latest_count or 3))
    delete_ids = [version.get("id") for version in sorted_versions[keep_latest_count:]]

    return delete_manuscript_version_ids(delete_ids)

st.markdown(
    """
    <style>
[data-testid="stMainBlockContainer"] {
    max-width: 1500px !important;
    padding-left: 2rem !important;
    padding-right: 2rem !important;
}
    div[role="radiogroup"] label {
        background-color: #ffffff !important;
        color: #000000 !important;
        border: 1px solid #dddddd !important;
        border-radius: 6px !important;
        padding: 8px 10px !important;
        margin-bottom: 8px !important;
        width: 100% !important;
        display: flex !important;
        align-items: center !important;
    }

    div[role="radiogroup"] label p {
        color: #000000 !important;
        margin: 0 !important;
    }

    div[role="radiogroup"] label span {
        color: #000000 !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ==========================================
# BACKEND INITIALIZATION
# ==========================================
if "backend" not in st.session_state:
    st.session_state.backend = GeminiWriterBackend()

load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")

supabase = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
# ==========================================
# SYSTEM STATE INITIALIZATION
# ==========================================
if "project_stage" not in st.session_state:
    st.session_state.project_stage = "Intake"

if "story_seed" not in st.session_state or not str(st.session_state.story_seed).strip():
    st.session_state.story_seed = (
        "An antique shop owner buys a heavy iron safe that cannot be opened, "
        "only to realize that every night, the ticking sounds coming from inside it get faster."
    )

if "project_title" not in st.session_state:
    st.session_state.project_title = ""

if "format_type" not in st.session_state:
    st.session_state.format_type = "Novel / Prose"

if "tone_style" not in st.session_state:
    st.session_state.tone_style = "Gritty Noir Mystery"

if "author_style" not in st.session_state:
    st.session_state.author_style = ""

if "sandbox_slates" not in st.session_state:
    st.session_state.sandbox_slates = None

if "chosen_option" not in st.session_state:
    st.session_state.chosen_option = None

if "chosen_option_text" not in st.session_state:
    st.session_state.chosen_option_text = None

if "plot_outline" not in st.session_state:
    st.session_state.plot_outline = None

if "current_draft_output" not in st.session_state:
    st.session_state.current_draft_output = ""

if "rolling_lore_book" not in st.session_state:
    st.session_state.rolling_lore_book = ""

if "completed_beats" not in st.session_state:
    st.session_state.completed_beats = []

if "scene_blocks" not in st.session_state:
    st.session_state.scene_blocks = {}

if "current_processing_beat" not in st.session_state:
    st.session_state.current_processing_beat = 1

if "manuscript_version" not in st.session_state:
    st.session_state.manuscript_version = 0

if "rewrite_applied_message" not in st.session_state:
    st.session_state.rewrite_applied_message = ""

if "rewritten_manuscript" not in st.session_state:
    st.session_state.rewritten_manuscript = ""

if "rewrite_preview_id" not in st.session_state:
    st.session_state.rewrite_preview_id = 0

if "rewrite_target_length" not in st.session_state:
    st.session_state.rewrite_target_length = "Keep Current Length"

if "writer_credits_remaining" not in st.session_state:
    st.session_state.writer_credits_remaining = 100

if "plot_outline_credit_pending" not in st.session_state:
    st.session_state.plot_outline_credit_pending = False

if "plot_outline_credit_cost" not in st.session_state:
    st.session_state.plot_outline_credit_cost = get_writer_credit_cost("plot_outline") if "get_writer_credit_cost" in globals() else 1

if "charged_scene_blocks" not in st.session_state:
    st.session_state.charged_scene_blocks = []

if "export_view" not in st.session_state:
    st.session_state.export_view = False

if "project_saved_notice" not in st.session_state:
    st.session_state.project_saved_notice = False

def count_words(text):
    return len(re.findall(r"\b\w+\b", str(text or "")))


def short_date_label(dt_value=None):
    if not dt_value:
        dt = datetime.now()
    else:
        try:
            dt = datetime.fromisoformat(str(dt_value).replace("Z", "+00:00"))
        except Exception:
            dt = datetime.now()
    return f"{dt.month}-{dt.day}-{str(dt.year)[-2:]}"


def get_project_title():
    return str(st.session_state.get("project_title", "") or "").strip()


def get_project_title_for_display():
    title = get_project_title()
    return title if title else "Untitled Project"


def make_version_display_label(version):
    version_number = version.get("version_number") or 1
    project_title = str(
        version.get("project_title")
        or version.get("title")
        or get_project_title_for_display()
    ).strip()

    version_label = str(version.get("version_label") or "Saved Version").strip()
    created_at = version.get("created_at")
    date_label = short_date_label(created_at)
    word_count = count_words(version.get("manuscript_content", ""))

    return f"{project_title}\n{date_label} — V{version_number} — {word_count} words"

def full_date_label(dt_value=None):
    if not dt_value:
        dt = datetime.now()
    else:
        try:
            dt = datetime.fromisoformat(str(dt_value).replace("Z", "+00:00"))
        except Exception:
            dt = datetime.now()
    return f"{dt.month:02d}/{dt.day:02d}/{dt.year}"


def get_saved_project_sessions_for_dropdown():
    user_id = get_current_user_id()

    query = (
        supabase.table("writersblock_project_drafts")
        .select("id, project_id, draft_content, created_at, version_number")
        .eq("draft_type", "project_session")
        .order("created_at", desc=True)
    )

    if user_id:
        query = query.eq("user_id", user_id)

    result = query.execute()
    return result.data or []


def get_saved_projects_versions_dropdown_items():
    items = []

    for saved_row in get_saved_project_sessions_for_dropdown():
        try:
            snapshot = json.loads(saved_row.get("draft_content") or "{}")
        except Exception:
            snapshot = {}

        title = str(snapshot.get("project_title") or "Untitled Project").strip()
        created_at = saved_row.get("created_at") or snapshot.get("saved_at")
        version_number = saved_row.get("version_number") or int(snapshot.get("manuscript_version") or 1)

        label = f"{title} — {full_date_label(created_at)} — V{version_number}"
        items.append({
            "label": label,
            "type": "project_session",
            "row": saved_row,
        })

    for version in get_manuscript_versions_for_current_project():
        project_join = version.get("writersblock_projects") or {}
        if isinstance(project_join, dict):
            version["project_title"] = project_join.get("project_title")

        title = str(
            version.get("project_title")
            or version.get("title")
            or get_project_title_for_display()
        ).strip()

        version_number = version.get("version_number") or 1
        created_at = version.get("created_at")
        label = f"{title} — {full_date_label(created_at)} — V{version_number}"

        items.append({
            "label": label,
            "type": "manuscript_version",
            "row": version,
        })

    return items


def get_current_plan_tier():
    if st.session_state.get("subscriber_mode"):
        return (st.session_state.get("plan_tier") or "").upper()
    if st.session_state.get("demo_mode"):
        return "DEMO"
    return ""


def plan_allows(required_stage: str):
    plan = get_current_plan_tier()

    plan_access = {
        "DEMO": ["intake", "sandbox", "plot", "drafting", "safety", "revision", "rewrite"],
        "STRUCTURE": ["intake", "sandbox", "plot"],
        "DRAFTING": ["intake", "sandbox", "plot", "drafting", "safety"],
        "PRODUCTION": ["intake", "sandbox", "plot", "drafting", "safety", "revision", "rewrite"],
    }

    return required_stage in plan_access.get(plan, [])


def access_is_active():
    return bool(st.session_state.get("subscriber_mode") or st.session_state.get("demo_mode"))


def require_active_access(required_stage=None):
    if not access_is_active():
        st.error("Access is not activated. Open the Access panel and enter a demo code or paid subscriber email before using this tool.")
        st.stop()

    if required_stage and not plan_allows(required_stage):
        st.error("Your current access level does not include this tool.")
        st.stop()

def get_writer_credit_cost(action_name, target_length="Keep Current Length"):
    if action_name == "sandbox":
        return 1
    if action_name == "plot_outline":
        return 1
    if action_name == "draft_scene":
        return 1
    if action_name == "revision":
        return 3
    if action_name == "rewrite_scene":
        return 2
    if action_name == "rewrite_full":
        if target_length == "Expand to 400 Words":
            return 3
        if target_length == "Shorten for Video / Script Adaptation":
            return 2
        return 2
    return 0

def show_writer_credit_warning(action_name, target_length="Keep Current Length"):
    cost = get_writer_credit_cost(action_name, target_length)

    if cost <= 0:
        return True

    remaining = st.session_state.get("writer_credits_remaining", 0)

    st.warning(
        f"⚠️ This action may use {cost} Writer Credits. "
        f"You currently have {remaining} credits remaining."
    )

    if remaining < cost:
        st.error("You do not have enough Writer Credits remaining for this action.")
        return False

    confirm_key = f"confirm_credit_use_{action_name}_{target_length}"

    return st.checkbox(
        f"I understand this action may use {cost} Writer Credits.",
        key=confirm_key
    )

def ai_output_is_usable(output):
    if output is None:
        return False

    text_output = str(output).strip()
    if not text_output:
        return False

    blocked_phrases = [
        "fallback output",
        "backend error:",
        "backend issue:",
        "returned an empty response",
        "did not return usable text",
        "local fallback",
        "fallback beat",
        "permission_denied",
        "403 permission",
        "403",
        "project has been denied access",
        "api issue detected",
    ]

    lowered = text_output.lower()
    return not any(phrase in lowered for phrase in blocked_phrases)

def get_writer_credit_balance():
    return int(st.session_state.get("writer_credits_remaining", 0) or 0)


def require_writer_credits(action_name, target_length="Keep Current Length"):
    cost = int(get_writer_credit_cost(action_name, target_length))

    if get_writer_credit_balance() < cost:
        st.error("Not enough Writer Credits for this action.")
        st.stop()

    return cost


def charge_writer_credits(action_name, output=None, target_length="Keep Current Length", require_usable_output=True):
    cost = int(get_writer_credit_cost(action_name, target_length))

    if cost <= 0:
        return False

    if require_usable_output and output is not None and not ai_output_is_usable(output):
        return False

    if get_writer_credit_balance() < cost:
        st.error("Not enough Writer Credits for this action.")
        st.stop()

    st.session_state.writer_credits_remaining = get_writer_credit_balance() - cost
    return True


def sandbox_output_is_usable(sandbox_output):
    return bool(
        sandbox_output
        and sandbox_output.get("Option A")
        and sandbox_output.get("Option B")
        and sandbox_output.get("Option C")
        and not sandbox_output.get("error_message")
    )


def get_demo_user_by_code(access_code: str):
    result = supabase.table("writersblock_users").select("*").eq("demo_code", access_code).limit(1).execute()
    if result.data and len(result.data) > 0:
        return result.data[0]
    return None

def get_subscriber_user_by_email(email: str):
    clean_email = (email or "").strip().lower()
    if not clean_email:
        return None

    result = (
        supabase.table("writersblock_users")
        .select("*")
        .eq("email", clean_email)
        .limit(1)
        .execute()
    )

    if result.data and len(result.data) > 0:
        return result.data[0]

    return None

def is_subscriber_access_valid(user_row: dict):
    if not user_row:
        return False, "Subscriber email not found."

    if not user_row.get("email_verified", False):
        return False, "Subscriber email is not verified."

    if not user_row.get("is_active", False):
        return False, "Subscriber account is not active."

    subscription_status = (user_row.get("subscription_status") or "").lower()
    if subscription_status not in ["active", "trialing"]:
        return False, "Subscriber plan is not active."

    plan_tier = user_row.get("plan_tier")
    if not plan_tier:
        return False, "No subscriber plan is assigned."

    return True, ""

def is_demo_access_valid(user_row: dict):
    if not user_row:
        return False, "Invalid demo code."

    if not user_row.get("demo_active", True):
        return False, "Demo access is inactive."

    used = int(user_row.get("demo_tokens_used") or 0)
    limit = int(user_row.get("demo_token_limit") or 5000)

    if used >= limit:
        return False, "Demo token limit reached."

    expires_at = user_row.get("demo_expires_at")
    if expires_at:
        expires_dt = datetime.fromisoformat(expires_at.replace("Z", "+00:00"))
        if datetime.now(timezone.utc) > expires_dt:
            return False, "Demo access expired."

    return True, ""


def get_demo_tokens_remaining(user_row: dict):
    used = int(user_row.get("demo_tokens_used") or 0)
    limit = int(user_row.get("demo_token_limit") or 5000)
    return max(0, limit - used)


def extract_total_tokens(response):
    usage = getattr(response, "usage_metadata", None)
    if not usage:
        return 0

    total = getattr(usage, "total_token_count", None)
    if total is not None:
        return int(total)

    prompt_tokens = getattr(usage, "prompt_token_count", 0) or 0
    candidate_tokens = getattr(usage, "candidates_token_count", 0) or 0
    return int(prompt_tokens) + int(candidate_tokens)


def add_demo_token_usage(user_id: str, token_count: int):
    current = supabase.table("writersblock_users").select("demo_tokens_used").eq("id", user_id).limit(1).execute()
    current_used = 0
    if current.data and len(current.data) > 0:
        current_used = int(current.data[0].get("demo_tokens_used") or 0)

    new_total = current_used + int(token_count)

    supabase.table("writersblock_users").update({
        "demo_tokens_used": new_total
    }).eq("id", user_id).execute()
def charge_ai_usage_from_text(text_output):
    if not st.session_state.get("demo_mode"):
        return

    user_row = refresh_demo_user_session()

    if not user_row:
        return

    tokens_used_now = max(1, len(str(text_output).split()))
    add_demo_token_usage(user_row["id"], tokens_used_now)

def refresh_demo_user_session():
    code = st.session_state.get("demo_access_code", "").strip()
    if not code:
        return None

    user_row = get_demo_user_by_code(code)
    st.session_state.demo_user_row = user_row
    return user_row
# ==========================================
# TITLE
# ==========================================
st.title("📝 WritersBlock Studios")
st.markdown("---")

# =====================================
# SIDEBAR STATUS / NAVIGATION / ACCESS
# =====================================

with st.sidebar.expander("Access", expanded=False):
    st.markdown(
        """
        <div style="text-align:center; margin-bottom:0.65rem;">
            <a href="https://writersblockstudios.com" target="_blank">
                Get Subscriber Access
            </a>
        </div>
        <div style="text-align:center; font-size:0.78rem; margin-bottom:0.65rem; opacity:0.85;">
            Already have access? Enter your demo code or subscriber email.
        </div>
        """,
        unsafe_allow_html=True,
    )

    active_access = st.session_state.get("subscriber_mode") or st.session_state.get("demo_mode")

    if st.session_state.get("access_success_message"):
        st.success(st.session_state.access_success_message)

    if active_access:
        if st.session_state.get("subscriber_mode"):
            st.markdown("**Current Access:** Paid Subscriber")
            st.markdown(f"**Plan:** {str(st.session_state.get('plan_tier', 'UNKNOWN')).upper()}")
            if st.session_state.get("subscriber_email"):
                st.markdown(f"**Email:** {st.session_state.subscriber_email}")
        elif st.session_state.get("demo_mode"):
            st.markdown("**Current Access:** Demo")
            if st.session_state.get("demo_access_code"):
                st.markdown(f"**Demo Code:** {st.session_state.demo_access_code}")

        if st.button("Switch User / Logout", key="switch_user_logout", use_container_width=True):
            st.session_state.demo_mode = False
            st.session_state.subscriber_mode = False
            st.session_state.demo_access_code = ""
            st.session_state.demo_user_row = None
            st.session_state.subscriber_user_row = None
            st.session_state.subscriber_email = ""
            st.session_state.plan_tier = ""
            st.session_state.access_success_message = ""
            st.rerun()

    else:
        access_mode = st.radio(
            "Access Type",
            ["Demo Code", "Paid Subscriber Email"],
            key="access_mode",
        )

        if access_mode == "Demo Code":
            st.markdown(
                """
                <div style="text-align:center; font-size:0.82rem; margin:0.45rem 0 0.25rem 0;">
                    Demo Access Code
                </div>
                """,
                unsafe_allow_html=True,
            )

            demo_code_input = st.text_input(
                "Demo Access Code",
                key="demo_access_code_input",
                label_visibility="collapsed",
            )

            if st.button("Activate Access", key="activate_demo_access", use_container_width=True):
                code = demo_code_input.strip()
                user_row = get_demo_user_by_code(code)
                ok, msg = is_demo_access_valid(user_row)

                if ok:
                    st.session_state.demo_access_code = code
                    st.session_state.demo_user_row = user_row
                    st.session_state.demo_mode = True
                    st.session_state.subscriber_mode = False
                    st.session_state.subscriber_user_row = None
                    st.session_state.subscriber_email = ""
                    st.session_state.plan_tier = "DEMO"
                    st.session_state.access_success_message = "Demo access activated."
                    st.rerun()
                else:
                    st.error(msg)

        elif access_mode == "Paid Subscriber Email":
            st.markdown(
                """
                <div style="text-align:center; font-size:0.82rem; margin:0.45rem 0 0.25rem 0;">
                    Subscriber Email
                </div>
                """,
                unsafe_allow_html=True,
            )

            subscriber_email_input = st.text_input(
                "Subscriber Email",
                key="subscriber_email_input",
                label_visibility="collapsed",
            )

            if st.button("Activate Access", key="activate_subscriber_access", use_container_width=True):
                email = subscriber_email_input.strip().lower()
                user_row = get_subscriber_user_by_email(email)
                ok, msg = is_subscriber_access_valid(user_row)

                if ok:
                    st.session_state.subscriber_user_row = user_row
                    st.session_state.subscriber_mode = True
                    st.session_state.demo_mode = False
                    st.session_state.demo_access_code = ""
                    st.session_state.demo_user_row = None
                    st.session_state.subscriber_email = email
                    st.session_state.plan_tier = user_row.get("plan_tier")
                    st.session_state.access_success_message = (
                        f"Paid subscriber access activated. Plan tier: "
                        f"{str(user_row.get('plan_tier', 'UNKNOWN')).upper()}"
                    )
                    st.rerun()
                else:
                    st.error(msg)

st.sidebar.markdown("---")

with st.sidebar.container(border=True):
    st.markdown("**📊 Project Status**")

    if st.session_state.get("subscriber_mode"):
        st.write("Access Mode: Paid Subscriber")
        st.write(f"Plan Tier: {st.session_state.get('plan_tier', 'Unknown')}")
    elif st.session_state.get("demo_mode"):
        st.write("Access Mode: Demo")
    else:
        st.write("Access Mode: Not Activated")

    st.write(f"Current Project: {get_project_title_for_display()}")
    st.write(f"Manuscript Version: {st.session_state.manuscript_version}")
    st.write(f"Current Manuscript Words: {count_words(st.session_state.current_draft_output)}")
    st.write(f"Writer Credits: {st.session_state.writer_credits_remaining}")

with st.sidebar.expander("Saved Projects / Versions", expanded=False):
    saved_items = get_saved_projects_versions_dropdown_items()

    if not saved_items:
        st.caption("No saved projects or versions found.")
    else:
        saved_labels = [item["label"] for item in saved_items]
        saved_lookup = {item["label"]: item for item in saved_items}

        selected_saved_label = st.selectbox(
            "Saved Projects / Versions:",
            saved_labels,
            key="sidebar_saved_projects_versions_selector",
        )

        selected_saved_item = saved_lookup.get(selected_saved_label)

        if st.button("Load Selected Saved Item", key="sidebar_load_saved_item", use_container_width=True):
            if selected_saved_item and selected_saved_item["type"] == "project_session":
                if load_saved_project_session_to_state(selected_saved_item["row"]):
                    st.success("Saved project loaded.")
                    st.rerun()

            elif selected_saved_item and selected_saved_item["type"] == "manuscript_version":
                selected_version = selected_saved_item["row"]
                st.session_state.current_project_id = selected_version.get("project_id")

                loaded_project_title = str(
                    selected_version.get("project_title")
                    or get_project_title_from_project_id(selected_version.get("project_id"))
                    or get_project_title()
                    or "Untitled Project"
                ).strip()

                st.session_state.project_title = loaded_project_title
                st.session_state.current_draft_output = str(selected_version.get("manuscript_content") or "")
                st.session_state.manuscript_version = int(selected_version.get("version_number") or 0)
                st.session_state.rewritten_manuscript = ""
                st.session_state.rewrite_applied_message = ""
                st.session_state.project_stage = "Isolation Guardrail"
                st.session_state.export_view = True
                st.success("Saved manuscript version loaded.")
                st.rerun()

        if selected_saved_item and selected_saved_item["type"] == "manuscript_version":
            try:
                selected_version = selected_saved_item["row"]
                selected_saved_text = str(selected_version.get("manuscript_content") or "")

                saved_doc = Document()
                saved_doc.add_heading(f"WritersBlock Studios Saved Manuscript - {selected_saved_label}", level=1)

                for paragraph in selected_saved_text.split("\n\n"):
                    clean_paragraph = paragraph.strip()
                    if clean_paragraph:
                        p = saved_doc.add_paragraph(clean_paragraph)
                        p.paragraph_format.space_after = Pt(8)
                        p.paragraph_format.line_spacing = 1.15

                saved_docx_buffer = BytesIO()
                saved_doc.save(saved_docx_buffer)

                st.download_button(
                    label="Download Selected Version (.docx)",
                    data=saved_docx_buffer.getvalue(),
                    file_name="saved_manuscript_version.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.caption(f"Download unavailable: {e}")

st.sidebar.markdown(
    """
    <div style="
        border:1px solid rgba(255,255,255,0.18);
        border-radius:10px;
        padding:0.55rem 0.55rem 0.48rem 0.55rem;
        margin:0.35rem 0 0.55rem 0;
        text-align:center;
    ">
        <h2 style="margin:0; padding:0; font-size:1.05rem; line-height:1.0;">
            The Writer Studios
        </h2>
        <p style="font-size:0.85rem; line-height:1.0; margin:0.18rem 0 0 0; padding:0;">
            Navigate your project through each creative production room.
        </p>
    </div>
    """,
    unsafe_allow_html=True,
)

can_navigate_to_sandbox = st.session_state.sandbox_slates is not None
can_navigate_to_plot = st.session_state.chosen_option is not None
can_navigate_to_drafting = st.session_state.plot_outline is not None
can_navigate_to_guardrail = bool(st.session_state.current_draft_output.strip())
can_navigate_to_revision = bool(st.session_state.current_draft_output.strip())
can_navigate_to_rewrite = bool(st.session_state.current_draft_output.strip())

can_use_sandbox = plan_allows("sandbox")
can_use_plot = plan_allows("plot")
can_use_drafting = plan_allows("drafting")
can_use_safety = plan_allows("safety")
can_use_revision = plan_allows("revision")
can_use_rewrite = plan_allows("rewrite")

if st.session_state.project_stage == "Intake":
    st.sidebar.button("🔵 1. Concept Development (Active)", key="nav_intake_active", use_container_width=True)
else:
    if st.sidebar.button("⚪ 1. Concept Development", key="nav_intake", use_container_width=True):
        st.session_state.project_stage = "Intake"
        st.rerun()

if st.session_state.project_stage == "Sandbox":
    st.sidebar.button("🔵 2. Scripting (Active)", key="nav_sandbox_active", use_container_width=True)
elif can_navigate_to_sandbox and can_use_sandbox:
    if st.sidebar.button("⚪ 2. Scripting", key="nav_sandbox", use_container_width=True):
        st.session_state.project_stage = "Sandbox"
        st.rerun()
else:
    st.sidebar.button("🔒 2. Scripting (Locked)", key="nav_sandbox_locked", disabled=True, use_container_width=True)

if st.session_state.project_stage == "Plot Architect":
    st.sidebar.button("🔵 3. Plot Architecture (Active)", key="nav_plot_active", use_container_width=True)
elif can_navigate_to_plot and can_use_plot:
    if st.sidebar.button("⚪ 3. Plot Architecture", key="nav_plot", use_container_width=True):
        st.session_state.project_stage = "Plot Architect"
        st.rerun()
else:
    st.sidebar.button("🔒 3. Plot Architecture (Locked)", key="nav_plot_locked", disabled=True, use_container_width=True)

if st.session_state.project_stage == "Drafting Engine":
    st.sidebar.button("🔵 4. Creative Draft (Active)", key="nav_draft_active", use_container_width=True)
elif can_navigate_to_drafting and can_use_drafting:
    if st.sidebar.button("⚪ 4. Creative Draft", key="nav_draft", use_container_width=True):
        st.session_state.project_stage = "Drafting Engine"
        st.rerun()
else:
    st.sidebar.button("🔒 4. Creative Draft (Locked)", key="nav_draft_locked", disabled=True, use_container_width=True)

if st.session_state.project_stage == "Isolation Guardrail" and not st.session_state.get("export_view", False):
    st.sidebar.button("🔵 5. Script Lock (Active)", key="nav_guard_active", use_container_width=True)
elif can_navigate_to_guardrail and can_use_safety:
    if st.sidebar.button("⚪ 5. Script Lock", key="nav_guard", use_container_width=True):
        st.session_state.project_stage = "Isolation Guardrail"
        st.session_state.export_view = False
        st.rerun()
else:
    st.sidebar.button("🔒 5. Script Lock (Locked)", key="nav_guard_locked", disabled=True, use_container_width=True)

if st.session_state.project_stage == "Revision Engine":
    st.sidebar.button("🔵 6. Script Polish (Active)", key="nav_revision_active", use_container_width=True)
elif can_navigate_to_revision and can_use_revision:
    if st.sidebar.button("⚪ 6. Script Polish", key="nav_revision", use_container_width=True):
        st.session_state.project_stage = "Revision Engine"
        st.rerun()
else:
    st.sidebar.button("🔒 6. Script Polish (Locked)", key="nav_revision_locked", disabled=True, use_container_width=True)

if st.session_state.project_stage == "Rewrite Engine":
    st.sidebar.button("🔵 7. Script Rewrite (Active)", key="nav_rewrite_active", use_container_width=True)
elif can_navigate_to_rewrite and can_use_rewrite:
    if st.sidebar.button("⚪ 7. Script Rewrite", key="nav_rewrite", use_container_width=True):
        st.session_state.project_stage = "Rewrite Engine"
        st.rerun()
else:
    st.sidebar.button("🔒 7. Script Rewrite (Locked)", key="nav_rewrite_locked", disabled=True, use_container_width=True)

if st.session_state.project_stage == "Isolation Guardrail" and st.session_state.get("export_view", False):
    st.sidebar.button("🔵 8. Manuscript Output (Active)", key="nav_export_active", use_container_width=True)
elif can_navigate_to_guardrail and can_use_safety:
    if st.sidebar.button("⚪ 8. Manuscript Output", key="nav_export", use_container_width=True):
        st.session_state.project_stage = "Isolation Guardrail"
        st.session_state.export_view = True
        st.rerun()
else:
    st.sidebar.button("🔒 8. Manuscript Output (Locked)", key="nav_export_locked", disabled=True, use_container_width=True)

st.sidebar.markdown("---")

with st.sidebar.container(border=True):
    st.markdown("**Current Project**")
    active_project_title = get_project_title()
    active_project_version = int(st.session_state.get("manuscript_version", 0) or 0) + 1
    active_saved_at = st.session_state.get("last_saved_project_at")

    if active_project_title:
        st.write(active_project_title)
        if active_saved_at:
            st.caption(f"{full_date_label(active_saved_at)} — V{active_project_version}")
        else:
            st.caption(f"Working Version: V{active_project_version}")
    else:
        st.caption("First name the project, then select Save Project")

if st.sidebar.button("💾 Save Project", key="save_current_project_session", use_container_width=True, type="secondary"):
    ensure_project_title_loaded()

    if not get_project_title():
        st.sidebar.warning("Project title could not be found. Return to Concept Development and name the project before saving.")
    else:
        result = save_current_project_session_to_supabase()
        if result:
            st.sidebar.success("Project saved.")
            st.rerun()

if can_navigate_to_guardrail and can_use_safety:
    if st.sidebar.button("📚 Saved Version Manager", key="sidebar_saved_version_manager_shortcut", use_container_width=True, type="secondary"):
        st.session_state.project_stage = "Isolation Guardrail"
        st.session_state.export_view = True
        st.rerun()
else:
    st.sidebar.button("🔒 Saved Version Manager", key="sidebar_saved_version_manager_locked", disabled=True, use_container_width=True, type="secondary")

if st.sidebar.button("⚠️ Start New Project", key="global_reset", use_container_width=True, type="secondary"):
    st.session_state.project_stage = "Intake"
    st.session_state.project_title = ""
    st.session_state.story_seed = (
        "An antique shop owner buys a heavy iron safe that cannot be opened, "
        "only to realize that every night, the ticking sounds coming from inside it get faster."
    )
    st.session_state.format_type = "Novel / Prose"
    st.session_state.tone_style = "Gritty Noir Mystery"
    st.session_state.author_style = ""
    st.session_state.sandbox_slates = None
    st.session_state.chosen_option = None
    st.session_state.chosen_option_text = None
    st.session_state.plot_outline = None
    st.session_state.current_draft_output = ""
    st.session_state.rolling_lore_book = ""
    st.session_state.completed_beats = []
    st.session_state.scene_blocks = {}
    st.session_state.current_processing_beat = 1
    st.session_state.revision_report = ""
    st.session_state.rewritten_manuscript = ""
    st.session_state.manuscript_version = 0
    st.session_state.rewrite_applied_message = ""
    st.session_state.current_project_id = None
    st.session_state.writer_credits_remaining = 100
    st.session_state.plot_outline_credit_pending = False
    st.session_state.plot_outline_credit_cost = 1
    st.session_state.charged_scene_blocks = []
    st.session_state.revision_report = ""
    st.session_state.export_view = False
    st.session_state.project_saved_notice = False
    st.session_state.rewrite_preview_id = 0
    st.session_state.rewrite_target_length = "Keep Current Length"
    st.rerun()

# ==========================================
# PHASE 1: CONCEPT DEVELOPMENT
# ==========================================
SPARK_BANK = {
    "YouTube Script (Long-form)": [
        "A deep dive into a bizarre real-world historical event where an entire village vanished for 48 hours, leaving behind identical pocket watches set to the wrong time.",
        "An investigation into a hidden internet mystery: an anonymous creator who has been uploading 1-second videos of the exact same empty hallway every single day since 2011.",
        "The psychological breakdown of the bystander effect online — how social media algorithms isolate users by feeding them custom alternate realities.",
    ],
    "Short-Form Video (TikTok/Shorts)": [
        "POV: You are a temporal detective from 2085, and you just realized your target is sitting directly across from you on the subway right now.",
        "The terrifying reason why every map from the 1600s had the exact same island marked DO NOT SAIL HERE — and why it just reappeared on modern satellite feeds.",
        "A psychological trick that reveals someone is lying based on the direction they glance when you say their middle name.",
    ],
    "Short Story": [
        "An antique shop owner buys a heavy iron safe that cannot be opened, only to realize that every night, the ticking sounds coming from inside it get faster.",
        "In a city where memories can be bought and sold like currency, a detective is hired to find a beautiful memory that someone was murdered to forget.",
        "A lighthouse keeper notices that the light is not reflecting off the ocean anymore — instead, something beneath the water is shining a different light back up at him.",
    ],
        "Audio Short Story — 15–30 min": [
        "A lighthouse keeper records one final message before the fog rolls in and the lanterns begin answering back.",
        "An elderly train conductor tells a passenger the same bedtime story every night, only to discover the story is changing.",
        "A woman staying alone in a quiet seaside inn hears a piano playing downstairs after midnight, though the inn has no piano.",
    ],
    "Screenplay Outline": [
        "A disgraced detective is forced to pair up with an AI android to solve a string of cyber-crimes, only to realize the android's code matches the handwriting of his missing daughter.",
        "A group of astronauts on a deep-space research vessel wake up from hyper-sleep to find an extra stasis pod on board that was not there when they launched.",
        "A forensic accountant auditing a massive tech corporation discovers a secret payroll line for an employee who has been paid continuously since 1895.",
    ],
    "Full Novel Blueprint": [
        "In a dystopian future where the human brain is permanently connected to a cloud network, a glitch causes one man to hear the internal monologues of everyone within a 1-mile radius.",
        "A young archeologist unearths a pristine modern digital watch buried inside an untouched 3,000-year-old Egyptian pharaoh's tomb.",
        "An elite hacker discovers a hidden string of code embedded in the human genome that acts as an IP address — and someone just tried to ping it from deep space.",
    ],
    "Children's Picture Book (32-Page Layout)": [
        "A tiny, brave field mouse finds a shiny fallen star in a meadow and journeys across the forest to toss it back into the night sky.",
        "A little boy discovers that his shadow does not want to copy him anymore — it wants to dance, explore, and teach him bravery.",
        "Barnaby the Bear is terrified of the dark until a wise old owl shows him that night is a soft velvet blanket embroidered with stars.",
    ],
}


if st.session_state.project_stage == "Intake":
    st.header("🗺️ Step 1: Concept Development")
    st.markdown("##### Build the foundation for your story, script, or manuscript.")

    st.session_state.project_title = st.text_input(
        "Project / Story Title:",
        value=st.session_state.get("project_title", ""),
        placeholder="Example: It's Not Safe",
        key="project_title_input",
    ).strip()

    format_options = list(SPARK_BANK.keys())

    st.session_state.format_type = st.selectbox(
        "Target Layout Format:",
        format_options,
        index=format_options.index(st.session_state.format_type)
        if st.session_state.format_type in format_options
        else format_options.index("Short Story"),
    )

    tone_options = [
        "Mind-Bending Sci-Fi",
        "Gritty Noir Mystery",
        "Children's: Scripture & Life Lessons",
        "Children's: Whimsical Animal Adventure",
        "Whimsical & Magical Fantasy",
        "Satirical & Sarcastic Comedy",
        "Suspenseful & Tense Thriller",
        "Melancholic & Nostalgic Prose",
        "Epic & Grandiose Lore",
    ]

    st.session_state.tone_style = st.selectbox(
        "Tone & Aesthetic Base:",
        tone_options,
        index=tone_options.index(st.session_state.tone_style)
        if st.session_state.tone_style in tone_options
        else tone_options.index("Gritty Noir Mystery"),
    )

    st.markdown("---")
    st.markdown("**🔴 Select a Starting Story Idea or Enter Your Own Plot Line:**")

    seed_options = SPARK_BANK[st.session_state.format_type] + ["-- Enter Custom Plot Line --"]

    current_seed = st.session_state.story_seed
    default_seed_index = 0
    if current_seed in seed_options:
        default_seed_index = seed_options.index(current_seed)
    elif current_seed not in SPARK_BANK[st.session_state.format_type]:
        default_seed_index = len(seed_options) - 1

    selected_seed = st.radio("Starting Story Ideas:", seed_options, index=default_seed_index)

    if selected_seed == "-- Enter Custom Plot Line --":
        st.session_state.story_seed = st.text_area(
            "Enter Custom Plot Line:",
            value=st.session_state.story_seed,
            height=140,
        )

        st.warning(
            "After typing your custom plot line, press CTRL + ENTER to apply it "
            "before clicking Generate Sandbox Paths."
        )
    else:
        st.session_state.story_seed = selected_seed

    st.session_state.author_style = st.text_input(
        "Target Author Style Matrix (Optional):",
        value=st.session_state.author_style,
        placeholder="e.g., Raymond Chandler, Philip K. Dick",
    )

    st.markdown("---")

    sandbox_credit_cost = get_writer_credit_cost("sandbox")
    sandbox_warning_ok = show_writer_credit_warning("sandbox")

    if st.button(
        f"Generate Sandbox Paths 🚀 — Uses {sandbox_credit_cost} Credit",
        key="generate_sandbox_paths",
        type="primary",
        use_container_width=True,
        disabled=not sandbox_warning_ok,
    ):
        require_active_access("sandbox")
        st.session_state.plot_outline = None
        st.session_state.current_draft_output = ""
        st.session_state.rewritten_manuscript = ""
        st.session_state.manuscript_text = ""
        st.session_state.current_processing_beat = 0
        st.session_state.total_beats = 0
        st.session_state.manuscript_version = 0

        if not get_project_title():
            st.error("Please name this project before continuing.")
            st.stop()

        if st.session_state.writer_credits_remaining < sandbox_credit_cost:
            st.error("Not enough Writer Credits for this action.")
            st.stop()

        if not st.session_state.story_seed.strip():
            st.error("Please enter a story concept before generating sandbox paths.")
            st.stop()

        with st.spinner("Processing voice matrices and structural paths..."):
            time.sleep(1)

            st.session_state.sandbox_slates = st.session_state.backend.generate_sandbox_slates(
                story_seed=st.session_state.story_seed,
                format_type=st.session_state.format_type,
                tone_style=st.session_state.tone_style,
                author_style=st.session_state.author_style,
            )

            if sandbox_output_is_usable(st.session_state.sandbox_slates):
                charge_writer_credits("sandbox", require_usable_output=False)

            save_project_to_supabase(
                project_title=get_project_title_for_display()[:80],
                project_type=st.session_state.format_type,
            )

            if st.session_state.get("demo_mode"):
                user_row = refresh_demo_user_session()
                if user_row:
                    tokens_used_now = max(1, len(str(st.session_state.sandbox_slates).split()))
                    add_demo_token_usage(user_row["id"], tokens_used_now)

            st.session_state.chosen_option = None
            st.session_state.chosen_option_text = None
            st.session_state.plot_outline = None
            st.session_state.current_draft_output = ""
            st.session_state.rolling_lore_book = ""
            st.session_state.completed_beats = []
            st.session_state.scene_blocks = {}
            st.session_state.current_processing_beat = 1
            st.session_state.project_saved_notice = True
            st.session_state.project_stage = "Sandbox"
            st.rerun()

# ======================================
# PHASE 2: SCRIPTING
# ======================================
elif st.session_state.project_stage == "Sandbox":
    st.header("🗺️ Step 2: Scripting")
    st.caption(f"Current Project: {get_project_title_for_display()}")
    st.markdown("##### Shape the first working script paths from your concept.")
    st.markdown("---")

    if st.session_state.get("project_saved_notice"):
        st.success("Project saved.")
        st.session_state.project_saved_notice = False

    sandbox_slates = st.session_state.get("sandbox_slates", {})

    if not sandbox_slates:
        st.warning("No story paths are available yet. Return to Concept Development and generate story paths first.")
    else:
        if sandbox_slates.get("is_fallback", False):
            st.warning("AI connection issue detected. Temporary local story paths are being shown.")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.subheader("Option A: Classic Structure")
            st.markdown("*A clear, traditional story path.*")
            st.info(sandbox_slates.get("Option A", ""))

            plot_credit_cost = get_writer_credit_cost("plot_outline")

            if st.button(f"Select Path A 🎯 — Uses {plot_credit_cost} Credit", key="btn_a", type="primary", use_container_width=True):
                require_active_access("plot")

                if st.session_state.writer_credits_remaining < plot_credit_cost:
                    st.error("Not enough Writer Credits for this action.")
                    st.stop()

                st.session_state.chosen_option = "Option A: Classic Structure"
                st.session_state.plot_outline_credit_pending = True
                st.session_state.plot_outline_credit_cost = plot_credit_cost
                st.session_state.chosen_option_text = sandbox_slates.get("Option A", "")

                latest_project = supabase.table("writersblock_projects").select("id").order("created_at", desc=True).limit(1).execute()

                if latest_project.data:
                    supabase.table("writersblock_project_drafts").insert({
                        "project_id": latest_project.data[0]["id"],
                        "user_id": get_current_user_id(),
                        "draft_title": "Option A: Classic Structure",
                        "draft_type": "sandbox_option",
                        "draft_content": st.session_state.chosen_option_text,
                        "version_number": 1,
                    }).execute()

                st.session_state.plot_outline = None
                st.session_state.project_stage = "Plot Architect"
                st.rerun()

        with col2:
            st.subheader("Option B: Twisted Perspective")
            st.markdown("*A stranger version with a hidden truth or reversal.*")
            st.info(sandbox_slates.get("Option B", ""))

            plot_credit_cost = get_writer_credit_cost("plot_outline")

            if st.button(f"Select Path B 🎯 — Uses {plot_credit_cost} Credit", key="btn_b", type="primary", use_container_width=True):
                require_active_access("plot")

                if st.session_state.writer_credits_remaining < plot_credit_cost:
                    st.error("Not enough Writer Credits for this action.")
                    st.stop()

                st.session_state.chosen_option = "Option B: Twisted Perspective"
                st.session_state.plot_outline_credit_pending = True
                st.session_state.plot_outline_credit_cost = plot_credit_cost
                st.session_state.chosen_option_text = sandbox_slates.get("Option B", "")

                latest_project = supabase.table("writersblock_projects").select("id").order("created_at", desc=True).limit(1).execute()

                if latest_project.data:
                    supabase.table("writersblock_project_drafts").insert({
                        "project_id": latest_project.data[0]["id"],
                        "user_id": get_current_user_id(),
                        "draft_title": "Option B: Twisted Perspective",
                        "draft_type": "sandbox_option",
                        "draft_content": st.session_state.chosen_option_text,
                        "version_number": 1,
                    }).execute()

                st.session_state.plot_outline = None
                st.session_state.project_stage = "Plot Architect"
                st.rerun()

        with col3:
            st.subheader("Option C: High-Impact Version")
            st.markdown("*A faster, more intense story path.*")
            st.info(sandbox_slates.get("Option C", ""))

            plot_credit_cost = get_writer_credit_cost("plot_outline")

            if st.button(f"Select Path C 🎯 — Uses {plot_credit_cost} Credit", key="btn_c", type="primary", use_container_width=True):
                require_active_access("plot")

                if st.session_state.writer_credits_remaining < plot_credit_cost:
                    st.error("Not enough Writer Credits for this action.")
                    st.stop()

                st.session_state.chosen_option = "Option C: High-Impact Version"
                st.session_state.plot_outline_credit_pending = True
                st.session_state.plot_outline_credit_cost = plot_credit_cost
                st.session_state.chosen_option_text = sandbox_slates.get("Option C", "")

                latest_project = supabase.table("writersblock_projects").select("id").order("created_at", desc=True).limit(1).execute()

                if latest_project.data:
                    supabase.table("writersblock_project_drafts").insert({
                        "project_id": latest_project.data[0]["id"],
                        "user_id": get_current_user_id(),
                        "draft_title": "Option C: High-Impact Version",
                        "draft_type": "sandbox_option",
                        "draft_content": st.session_state.chosen_option_text,
                        "version_number": 1,
                    }).execute()

                st.session_state.plot_outline = None
                st.session_state.project_stage = "Plot Architect"
                st.rerun()

# ==========================================
# PHASE 3: PLOT ARCHITECT ENGINE
# ==========================================
if st.session_state.project_stage == "Plot Architect":
    st.header("🏗️ Step 3: Plot Architecture")
    st.caption(f"Current Project: {get_project_title_for_display()}")
    st.markdown(f"**Selected Direction:** `{st.session_state.chosen_option}`")
    st.markdown("---")

    if st.session_state.plot_outline is None:
        plot_credit_cost = int(st.session_state.get("plot_outline_credit_cost", get_writer_credit_cost("plot_outline")))
        plot_warning_ok = show_writer_credit_warning("plot_outline")

        if st.button(
            f"Generate Plot Architecture 🧱 — Uses {plot_credit_cost} Credit",
            key="generate_plot_architecture",
            type="primary",
            use_container_width=True,
            disabled=not plot_warning_ok,
        ):
            if st.session_state.writer_credits_remaining < plot_credit_cost:
                st.error("Not enough Writer Credits for this action.")
                st.stop()

            with st.spinner("Locking down scene beats and style directives..."):
                time.sleep(1)
                st.session_state.plot_outline = st.session_state.backend.generate_plot_outline(
                    chosen_option=st.session_state.chosen_option,
                    option_text=st.session_state.chosen_option_text,
                    story_seed=st.session_state.story_seed,
                    format_type=st.session_state.format_type,
                    tone_style=st.session_state.tone_style,
                    author_style=st.session_state.author_style,
                )

            if ai_output_is_usable(st.session_state.plot_outline):
                charge_writer_credits("plot_outline", st.session_state.plot_outline)

            st.session_state.plot_outline_credit_pending = False
            st.rerun()
        else:
            st.stop()

    st.subheader("Review / Modify Architectural Output:")

    st.session_state.plot_outline = st.text_area(
        "Structured Scene Beats Layout Panel:",
        value=st.session_state.plot_outline,
        height=450,
        key="current_outline_view",
        label_visibility="collapsed",
    )

    st.markdown("---")

    if st.button("Lock Outline & Advance to Recursive Drafting 🚀", key="lock_outline_advance", type="primary", use_container_width=True):
        require_active_access("drafting")

        st.session_state.current_draft_output = ""
        st.session_state.rolling_lore_book = ""
        st.session_state.completed_beats = []
        st.session_state.scene_blocks = {}
        st.session_state.current_processing_beat = 1
        st.session_state.project_stage = "Drafting Engine"
        st.rerun()

# ======================================
# PHASE 4: RECURSIVE DRAFTING ENGINE
# ======================================
elif st.session_state.project_stage in ["Drafting", "Drafting Engine"]:
    require_active_access("drafting")

    st.header("📄 Step 4: Creative Draft")
    st.caption(f"Current Project: {get_project_title_for_display()}")
  

    if "current_draft_output" not in st.session_state:
        st.session_state.current_draft_output = ""

    if "rolling_lore_book" not in st.session_state:
        st.session_state.rolling_lore_book = ""

    if "current_processing_beat" not in st.session_state:
        st.session_state.current_processing_beat = 1

    plot_outline_text = st.session_state.get("plot_outline", "")

    filtered_beats = [
        beat.strip()
        for beat in str(plot_outline_text).split("===")
        if beat.strip()
    ]

    if not filtered_beats:
        if str(plot_outline_text).strip():
            filtered_beats = [str(plot_outline_text).strip()]
        else:
            filtered_beats = []

    total_beats = len(filtered_beats)

    if total_beats == 0:
        st.warning("No plot outline beats found. Return to Plot Architect and create an outline first.")

    else:
        if "scene_blocks" not in st.session_state:
            st.session_state.scene_blocks = {}

        saved_scene_numbers = [
            int(scene_number)
            for scene_number, scene_text in st.session_state.scene_blocks.items()
            if str(scene_text or "").strip()
        ]

        next_missing_scene = None
        for scene_number in range(1, total_beats + 1):
            if scene_number not in saved_scene_numbers:
                next_missing_scene = scene_number
                break

        if next_missing_scene is not None:
            st.session_state.current_processing_beat = next_missing_scene
        else:
            st.session_state.current_processing_beat = total_beats + 1

        col_left, col_right = st.columns([2, 1])

        with col_left:
            st.subheader("📄 Draft Output Canvas")

            latest_scene_number = max(1, int(st.session_state.get("current_processing_beat", 1) or 1) - 1)
            latest_scene_text = ""
            if isinstance(st.session_state.get("scene_blocks"), dict):
                latest_scene_text = st.session_state.scene_blocks.get(latest_scene_number, "")

            st.text_area(
                "Latest Scene Output:",
                value=latest_scene_text or st.session_state.current_draft_output,
                height=500,
                label_visibility="collapsed",
                disabled=True,
            )

            st.subheader("⚙️ Chapter / Scene Settings")

            chapter_title = st.text_input(
                "Chapter / Scene Title:",
                value=f"Scene Block {st.session_state.current_processing_beat}",
                key=f"chapter_title_{st.session_state.current_processing_beat}"
            )

            target_word_count = st.number_input(
                "Target Word Count Per Scene Block:",
                min_value=50,
                max_value=2500,
                value=1200,
                step=25,
                key=f"target_word_count_{st.session_state.current_processing_beat}"
            )

            pov_choice = st.selectbox(
                "Point of View:",
                ["Third Person Limited", "Third Person Omniscient", "First Person"],
                key=f"pov_choice_{st.session_state.current_processing_beat}"
            )

            tense_choice = st.selectbox(
                "Tense:",
                ["Past Tense", "Present Tense"],
                key=f"tense_choice_{st.session_state.current_processing_beat}"
            )

            scene_notes = st.text_area(
                "Scene Purpose / Special Instructions:",
                placeholder="Example: Make this scene slower, more suspenseful, with stronger sensory detail.",
                height=100,
                key=f"scene_notes_{st.session_state.current_processing_beat}"
            )

            st.markdown("---")

        if st.session_state.current_processing_beat <= total_beats:
            current_scene_number = st.session_state.current_processing_beat

            if current_scene_number in st.session_state.scene_blocks and str(st.session_state.scene_blocks.get(current_scene_number, "")).strip():
                st.info(f"Scene Block {current_scene_number} is already written.")
                st.stop()

            scene_credit_cost = get_writer_credit_cost("draft_scene")
            scene_warning_ok = show_writer_credit_warning("draft_scene")

            btn_label = (
                f"Execute Scene Block "
                f"{current_scene_number} of {total_beats} ⚡ — Uses {scene_credit_cost} Credit"
            )

            if st.button(
                btn_label,
                key=f"execute_scene_block_{current_scene_number}",
                type="primary",
                use_container_width=True,
                disabled=not scene_warning_ok,
            ):
                current_beat_data = filtered_beats[current_scene_number - 1]
                require_writer_credits("draft_scene")

                drafting_instruction_packet = f"""
CHAPTER / SCENE SETTINGS:
Chapter / Scene Title: {chapter_title}
Target Word Count: Approximately {target_word_count} words
Point of View: {pov_choice}
Tense: {tense_choice}
Scene Purpose / Special Instructions: {scene_notes}

CURRENT OUTLINE BEAT:
{current_beat_data}

IMPORTANT OUTPUT RULE:
Write ONLY this single scene block. Do not repeat earlier scenes. Do not summarize the full manuscript. Do not include manuscript output, prior scene text, headings, or any material outside this current scene block. End with a LORE UPDATE only if needed.
"""

                with st.spinner(f"Drafting scene block {current_scene_number}..."):
                        try:
                            if st.session_state.get("demo_mode"):
                                user_row = refresh_demo_user_session()
                                ok, msg = is_demo_access_valid(user_row)
                                if not ok:
                                    st.error(msg)
                                    st.stop()

                            raw_response = st.session_state.backend.execute_recursive_draft_chunk(
                                current_beat_text=drafting_instruction_packet,
                                rolling_lore=st.session_state.rolling_lore_book,
                                format_type=st.session_state.format_type,
                                tone_style=st.session_state.tone_style,
                                author_style=st.session_state.author_style,
                            )

                        except Exception as e:
                            st.error(f"Drafting failed. No scene text was saved. Backend error: {e}")
                            st.stop()

                        if not ai_output_is_usable(raw_response):
                            st.error("Drafting failed. Gemini did not return usable scene text. No scene text was saved and no Writer Credits were charged.")
                            st.stop()

                        charge_ai_usage_from_text(raw_response)

                        lore_marker_match = re.search(
                            r"===?\s*LORE\s+UPDATE\s*===?",
                            str(raw_response),
                            re.IGNORECASE,
                        )

                        if lore_marker_match:
                            start_idx, end_idx = lore_marker_match.span()
                            narrative_text = str(raw_response)[:start_idx].strip()
                            new_lore = str(raw_response)[end_idx:].strip()
                        else:
                            narrative_text = str(raw_response).strip()
                            new_lore = (
                                st.session_state.rolling_lore_book
                                if st.session_state.rolling_lore_book
                                else "- Active context state synchronized."
                            )

                        if narrative_text:
                            if "scene_blocks" not in st.session_state:
                                st.session_state.scene_blocks = {}

                            if current_scene_number in st.session_state.scene_blocks:
                                st.warning(f"Scene Block {current_scene_number} was already written. It was not added twice.")
                                st.stop()

                            st.session_state.scene_blocks[current_scene_number] = narrative_text

                            ordered_scene_texts = []
                            for scene_idx in sorted(st.session_state.scene_blocks.keys()):
                                scene_text = str(st.session_state.scene_blocks.get(scene_idx) or "").strip()
                                if scene_text:
                                    ordered_scene_texts.append(scene_text)

                            st.session_state.current_draft_output = "\n\n".join(ordered_scene_texts)

                        if new_lore:
                            st.session_state.rolling_lore_book = new_lore

                        if narrative_text:
                            if "charged_scene_blocks" not in st.session_state:
                                st.session_state.charged_scene_blocks = []

                            if current_scene_number not in st.session_state.charged_scene_blocks:
                                charge_writer_credits("draft_scene", narrative_text, require_usable_output=True)
                                st.session_state.charged_scene_blocks.append(current_scene_number)

                        # Save generated scene block to database.
                        latest_project = supabase.table("writersblock_projects").select("id").order("created_at", desc=True).limit(1).execute()

                        if latest_project.data and narrative_text:
                            supabase.table("writersblock_project_drafts").insert({
                                "project_id": latest_project.data[0]["id"],
                                "user_id": get_current_user_id(),
                                "draft_title": f"Scene Block {current_scene_number}",
                                "draft_type": "scene_draft",
                                "draft_content": narrative_text,
                                "version_number": current_scene_number
                            }).execute()

                        saved_scene_numbers = [
                            int(scene_number)
                            for scene_number, scene_text in st.session_state.scene_blocks.items()
                            if str(scene_text or "").strip()
                        ]

                        next_missing_scene = None
                        for scene_number in range(1, total_beats + 1):
                            if scene_number not in saved_scene_numbers:
                                next_missing_scene = scene_number
                                break

                        if next_missing_scene is None:
                            st.session_state.current_processing_beat = total_beats + 1
                            st.session_state.project_stage = "Isolation Guardrail"
                        else:
                            st.session_state.current_processing_beat = next_missing_scene

                        st.rerun()

        with col_right:
            st.subheader("🧠 Rolling Lore Memory")

            st.text_area(
                "Rolling Lore Memory:",
                value=st.session_state.rolling_lore_book,
                height=250,
                label_visibility="collapsed",
                disabled=True,
            )

            st.markdown("---")

            saved_scene_count = len([
                scene_text
                for scene_text in st.session_state.scene_blocks.values()
                if str(scene_text or "").strip()
            ])

            st.metric(
                "Current Scene Beat",
                f"{saved_scene_count} / {total_beats}"
            )

            st.markdown("---")

            for idx in range(1, total_beats + 1):
                scene_is_saved = (
                    idx in st.session_state.scene_blocks
                    and str(st.session_state.scene_blocks.get(idx, "")).strip()
                )

                if scene_is_saved:
                    st.markdown(f"✅ Scene Block {idx}: Written")
                elif idx == st.session_state.current_processing_beat:
                    st.markdown(f"🔷 Scene Block {idx}: Ready")
                else:
                    st.markdown(f"⚪ Scene Block {idx}: Pending")
# ======================================================================================
# PHASE 5: SCRIPT LOCK / MANUSCRIPT OUTPUT
# ======================================================================================
elif st.session_state.project_stage == "Isolation Guardrail":
    require_active_access("safety")

    if st.session_state.get("export_view", False):
        st.header("📤 Step 8: Manuscript Output")
        st.caption(f"Current Project: {get_project_title_for_display()}")
        st.markdown("##### Review, select, and download your manuscript version.")
    else:
        st.header("🛡️ Step 5: Script Lock")
        st.caption(f"Current Project: {get_project_title_for_display()}")
        st.markdown("##### Verify the final manuscript before output or script rewrite.")
    st.markdown("---")

    st.info("✅ Script Lock check complete. Manuscript is ready for output or script rewrite.")
    col_left, col_right = st.columns([2, 1])

    with col_left:
        st.subheader("📋 Manuscript Output")

        st.text_area(
            "Review Manuscript Output:",
            value=st.session_state.current_draft_output,
            height=500,
            key="final_manuscript_export_view",
        )

        try:
            download_options = {
                "Current Manuscript": st.session_state.current_draft_output
            }

            if st.session_state.get("rewritten_manuscript", "").strip():
                download_options["Script Rewrite Preview - Not Accepted Yet"] = st.session_state.rewritten_manuscript

            saved_versions_data = get_manuscript_versions_for_current_project()

            if saved_versions_data:
                for version in saved_versions_data:
                    version_number = version.get("version_number")
                    version_label = version.get("version_label") or "Saved Version"
                    label = f"Version {version_number} - {version_label}"
                    download_options[label] = version.get("manuscript_content", "")

            selected_download_label = st.selectbox(
                "Choose manuscript version to download:",
                list(download_options.keys()),
                key="safety_guardrail_download_version"
            )

            selected_download_text = download_options.get(selected_download_label, "")

            doc = Document()
            doc.add_heading(f"WritersBlock Studios Manuscript - {selected_download_label}", level=1)

            for paragraph in str(selected_download_text).split("\n\n"):
                clean_paragraph = paragraph.strip()
                if clean_paragraph:
                    p = doc.add_paragraph(clean_paragraph)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = 1.15

            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()

            safe_filename = (
                selected_download_label
                .lower()
                .replace(" ", "_")
                .replace("-", "_")
                .replace("/", "_")
            )

            st.download_button(
                label=f"Download {selected_download_label} (.docx)",
                data=docx_bytes,
                file_name=f"{safe_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )

            if can_use_rewrite:
                if st.button(
                    "Proceed to Script Rewrite 🛠️",
                    key="proceed_to_rewrite_from_export",
                    type="primary",
                    use_container_width=True,
                ):
                    st.session_state.project_stage = "Rewrite Engine"
                    st.rerun()
            else:
                st.button(
                    "🔒 Script Rewrite Locked",
                    key="rewrite_locked_from_export",
                    disabled=True,
                    use_container_width=True,
                )

        except Exception as e:
            st.error(f"Failed to generate Word document: {e}")

    with col_right:
        st.subheader("📊 Output Summary")

        compiled_scene_count = len([
            scene_text
            for scene_text in st.session_state.get("scene_blocks", {}).values()
            if str(scene_text or "").strip()
        ])

        if compiled_scene_count == 0 and str(st.session_state.get("current_draft_output", "")).strip():
            compiled_scene_count = max(int(st.session_state.get("current_processing_beat", 1) or 1) - 1, 1)

        st.metric(
            label="Script Blocks Compiled",
            value=f"{compiled_scene_count}",
        )

        st.markdown("**Project Memory Cache:**")

        st.text_area(
            "Saved Project Memory:",
            value=st.session_state.rolling_lore_book,
            height=200,
            key="final_lore_view",
            disabled=True,
        )

        st.write("---")
        st.subheader("🗂️ Saved Versions")

        saved_version_count = len(get_manuscript_versions_for_current_project())
        st.caption(f"Saved manuscript versions: {saved_version_count}")
        st.caption("Open Saved Version Manager below to review, restore, download, or delete one saved version at a time.")

        st.write("---")

        if st.button("🔄 Start New Project", use_container_width=True):
            st.session_state.project_stage = "Intake"
            st.session_state.sandbox_slates = None
            st.session_state.chosen_option = None
            st.session_state.chosen_option_text = None
            st.session_state.plot_outline = None
            st.session_state.current_draft_output = ""
            st.session_state.rolling_lore_book = ""
            st.session_state.completed_beats = []
            st.session_state.scene_blocks = {}
            st.session_state.current_processing_beat = 1
            st.session_state.rewritten_manuscript = ""
            st.session_state.rewrite_applied_message = ""
            st.session_state.manuscript_version = 0
            st.session_state.export_view = False
            st.rerun()

    if st.session_state.get("export_view", False):
        st.markdown("---")
        st.subheader("🗂️ Saved Version Manager")
        st.caption(
            "Review saved manuscript versions here. Nothing changes unless you choose Restore or Delete. "
            "Delete removes only the selected saved version after confirmation."
        )

        manager_versions = get_manuscript_versions_for_current_project()

        if not manager_versions:
            st.info("No saved manuscript versions found yet.")
        else:
            manager_items = []

            for version in manager_versions:
                project_join = version.get("writersblock_projects") or {}
                if isinstance(project_join, dict):
                    project_title_for_version = project_join.get("project_title")
                else:
                    project_title_for_version = None

                title = str(
                    project_title_for_version
                    or version.get("project_title")
                    or get_project_title_for_display()
                ).strip()

                version_number = int(version.get("version_number") or 1)
                created_at = version.get("created_at")
                word_count = count_words(version.get("manuscript_content", ""))
                label_text = str(version.get("version_label") or "Saved Version").strip()
                display_label = f"{title} — {full_date_label(created_at)} — V{version_number} — {word_count} words — {label_text}"

                manager_items.append({
                    "label": display_label,
                    "version": version,
                    "title": title,
                    "version_number": version_number,
                    "word_count": word_count,
                    "label_text": label_text,
                })

            manager_labels = [item["label"] for item in manager_items]
            manager_lookup = {item["label"]: item for item in manager_items}

            selected_manager_label = st.selectbox(
                "Saved versions:",
                manager_labels,
                key="saved_version_manager_selector",
            )

            selected_manager_item = manager_lookup.get(selected_manager_label)
            selected_manager_version = selected_manager_item["version"] if selected_manager_item else None
            selected_manager_text = str(selected_manager_version.get("manuscript_content") or "") if selected_manager_version else ""

            st.text_area(
                "Selected Version Review:",
                value=selected_manager_text,
                height=520,
                key="saved_version_manager_review_box",
                disabled=True,
            )

            col_download, col_restore, col_delete = st.columns(3)

            with col_download:
                try:
                    version_doc = Document()
                    version_doc.add_heading(f"WritersBlock Studios Manuscript - {selected_manager_item['label']}", level=1)

                    for paragraph in selected_manager_text.split("\n\n"):
                        clean_paragraph = paragraph.strip()
                        if clean_paragraph:
                            p = version_doc.add_paragraph(clean_paragraph)
                            p.paragraph_format.space_after = Pt(8)
                            p.paragraph_format.line_spacing = 1.15

                    version_buffer = BytesIO()
                    version_doc.save(version_buffer)

                    st.download_button(
                        label="Download Selected Version",
                        data=version_buffer.getvalue(),
                        file_name=f"saved_version_v{selected_manager_item['version_number']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"Selected version download unavailable: {e}")

            with col_restore:
                restore_confirmed = st.checkbox(
                    "Confirm restore",
                    key="confirm_restore_selected_version",
                )

                if st.button(
                    "Restore Selected Version",
                    key="restore_selected_saved_version",
                    disabled=not restore_confirmed,
                    use_container_width=True,
                ):
                    st.session_state.current_draft_output = selected_manager_text
                    st.session_state.current_project_id = selected_manager_version.get("project_id")
                    st.session_state.project_title = selected_manager_item.get("title") or get_project_title_for_display()
                    st.session_state.manuscript_version = int(selected_manager_version.get("version_number") or st.session_state.get("manuscript_version", 0) or 0)
                    st.session_state.rewritten_manuscript = ""
                    st.session_state.rewrite_applied_message = "Selected saved version restored as the current manuscript."
                    save_current_project_session_to_supabase()
                    st.success("Selected saved version restored as the current manuscript.")
                    st.rerun()

            with col_delete:
                delete_confirmed = st.checkbox(
                    "I understand this will permanently delete the selected saved version.",
                    key="confirm_delete_selected_version",
                )

                if st.button(
                    "Delete Selected Version",
                    key="delete_selected_saved_version",
                    disabled=not delete_confirmed,
                    use_container_width=True,
                ):
                    delete_manuscript_version_ids([selected_manager_version.get("id")])
                    st.success("Selected saved version deleted.")
                    st.rerun()
# ======================================================================================
# PHASE 6: SCRIPT POLISH
# ======================================================================================
if st.session_state.project_stage == "Revision Engine":
    require_active_access("revision")

    st.header("🧠 Step 6: Script Polish")
    st.caption(f"Current Project: {get_project_title_for_display()}")
    st.markdown("##### Review the manuscript for continuity, pacing, tone, and polish opportunities.")
    st.markdown("---")

    if not st.session_state.current_draft_output.strip():
        st.warning("No manuscript draft found. Return to Creative Draft and complete drafting first.")
    else:
        col_left, col_right = st.columns([2, 1])

        with col_left:
            st.subheader("📄 Script Under Review")

            st.text_area(
                "Current Manuscript:",
                value=st.session_state.current_draft_output,
                height=500,
                label_visibility="collapsed",
                disabled=True,
            )

        with col_right:
            st.subheader("🔍 Polish Tools")

            revision_mode = st.selectbox(
                "Choose Polish Pass:",
                [
                    "Continuity Check",
                    "Pacing Check",
                    "Tone & Style Check",
                    "Character Consistency Check",
                    "Scene Expansion Suggestions",
                    "Final Polish Pass",
                ],
            )

            revision_credit_cost = get_writer_credit_cost("revision")

            if st.button(f"Run Script Polish 🧠 — Uses {revision_credit_cost} Credit", key="run_revision_pass", type="primary", use_container_width=True):
                require_writer_credits("revision")
                with st.spinner("Running polish analysis..."):
                    revision_report = st.session_state.backend.run_revision_pass(
                        manuscript_text=st.session_state.current_draft_output,
                        revision_mode=revision_mode,
                        tone_style=st.session_state.tone_style,
                        author_style=st.session_state.author_style,
                    )

                    st.session_state.revision_report = revision_report
                    if ai_output_is_usable(revision_report):
                        charge_writer_credits("revision", revision_report)
                    st.rerun()

            st.markdown("---")

            if "revision_report" not in st.session_state:
                st.session_state.revision_report = ""

            st.text_area(
                "Polish Report:",
                value=st.session_state.revision_report,
                height=300,
                label_visibility="collapsed",
                disabled=True,
            )

            if can_use_rewrite:
                if st.button("Proceed to Script Rewrite 🛠️", key="proceed_to_rewrite_from_revision", type="primary", use_container_width=True):
                    st.session_state.project_stage = "Rewrite Engine"
                    st.rerun()
            else:
                st.button("🔒 Script Rewrite Locked", key="rewrite_locked_from_revision", disabled=True, use_container_width=True)
# ============================================================
# PHASE 7: SCRIPT REWRITE
# ============================================================
elif st.session_state.project_stage == "Rewrite Engine":
    require_active_access("rewrite")

    st.header("🛠️ Step 7: Script Rewrite")
    st.caption(f"Current Project: {get_project_title_for_display()}")
    st.markdown("##### Script Rewrite uses the current live manuscript only. Saved versions are available for download, but are not used as the script rewrite source.")
    st.markdown("---")

    if "rewritten_manuscript" not in st.session_state:
        st.session_state.rewritten_manuscript = ""

    if "rewrite_preview_id" not in st.session_state:
        st.session_state.rewrite_preview_id = 0

    if "rewrite_target_length" not in st.session_state:
        st.session_state.rewrite_target_length = "Keep Current Length"

    manuscript_text = st.session_state.get("current_draft_output", "")

    def split_text_into_scene_chunks(full_text, total_scenes=4):
        full_text = str(full_text or "").strip()
        if not full_text:
            return []

        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", full_text) if p.strip()]
        if not paragraphs:
            return [full_text]

        total_scenes = max(1, int(total_scenes or 1))
        chunk_size = max(1, (len(paragraphs) + total_scenes - 1) // total_scenes)

        chunks = []
        for i in range(total_scenes):
            start_idx = i * chunk_size
            end_idx = min(len(paragraphs), (i + 1) * chunk_size)
            if start_idx < len(paragraphs):
                chunks.append("\n\n".join(paragraphs[start_idx:end_idx]).strip())

        return chunks

    def get_scene_beat_text(full_text, beat_number):
        scene_blocks = st.session_state.get("scene_blocks", {})
        stored_scene = scene_blocks.get(int(beat_number)) if isinstance(scene_blocks, dict) else None
        if stored_scene and str(stored_scene).strip():
            return str(stored_scene).strip()

        chunks = split_text_into_scene_chunks(full_text, 4)
        idx = int(beat_number) - 1
        if 0 <= idx < len(chunks):
            return chunks[idx]
        return ""

    def replace_scene_beat_text(full_text, beat_number, replacement_text):
        full_text = str(full_text or "")
        replacement_text = str(replacement_text or "").strip()

        scene_blocks = st.session_state.get("scene_blocks", {})
        stored_scene = None

        if isinstance(scene_blocks, dict):
            stored_scene = scene_blocks.get(int(beat_number))

        if stored_scene and str(stored_scene).strip() and str(stored_scene).strip() in full_text:
            st.session_state.scene_blocks[int(beat_number)] = replacement_text
            return full_text.replace(str(stored_scene).strip(), replacement_text, 1)

        current_scene_text = get_scene_beat_text(full_text, beat_number)
        if current_scene_text and current_scene_text in full_text:
            if isinstance(scene_blocks, dict):
                st.session_state.scene_blocks[int(beat_number)] = replacement_text
            return full_text.replace(current_scene_text, replacement_text, 1)

        return full_text

    def strip_internal_notices(text):
        text = str(text or "")
        text = re.sub(r"\n?\[LOCAL NODE 8 FALLBACK NOTE:.*?\]", "", text, flags=re.DOTALL)
        text = re.sub(r"\n?\[LOCAL .*?FALLBACK NOTE:.*?\]", "", text, flags=re.DOTALL)
        text = re.sub(r"\n?NODE 8 TEST OUTPUT.*", "", text, flags=re.DOTALL)
        return text.strip()

    if not manuscript_text.strip():
        st.warning("No manuscript draft found. Return to Creative Draft and complete drafting first.")
    else:
        col_left, col_right = st.columns([2, 1])

        with col_right:
            st.subheader("🛠️ Script Rewrite Controls")

            rewrite_mode = st.selectbox(
                "1) Script Rewrite Mode:",
                [
                    "Creative Script Rewrite / Improve Scene",
                    "Direct Script Change / Obey My Instructions Exactly",
                    "Polish Only / Do Not Change Story",
                ],
                key="rewrite_mode_select",
            )

            scene_choice = st.selectbox(
                "2) Script Rewrite Location:",
                [
                    "Full Manuscript",
                    "Scene Beat 1",
                    "Scene Beat 2",
                    "Scene Beat 3",
                    "Scene Beat 4",
                ],
                key="rewrite_scene_choice",
            )

            rewrite_target_length = st.selectbox(
                "3) Target Length:",
                [
                    "Keep Current Length",
                    "Shorten for Video / Script Adaptation",
                    "Expand to 400 Words",
                    "Custom Word Count",
                ],
                key="rewrite_target_length",
            )

            custom_target_word_count = None

            if rewrite_target_length == "Custom Word Count":
                custom_target_word_count = st.number_input(
                    "Custom Target Word Count:",
                    min_value=50,
                    max_value=2500,
                    value=min(1200, max(50, count_words(st.session_state.current_draft_output))),
                    step=25,
                    key="custom_target_word_count",
                )

            custom_rewrite_instruction = st.text_area(
                "4) Specific Script Rewrite Direction - Optional:",
                value="",
                height=160,
                placeholder=(
                    "Example: Keep the same ending, but expand the story with stronger atmosphere, "
                    "more antique shop detail, more dread, and more of Elias's physical decline."
                ),
                key="custom_rewrite_instruction",
            )

            st.markdown("---")
            st.info(
                "Script Rewrite uses only the current live manuscript shown on this page. "
                "Saved versions are not used as the script rewrite source."
            )

        if scene_choice == "Full Manuscript":
            selected_source_text = manuscript_text
            rewrite_credit_action = "rewrite_full"
        else:
            beat_number = int(scene_choice.replace("Scene Beat ", ""))
            selected_source_text = get_scene_beat_text(manuscript_text, beat_number)
            rewrite_credit_action = "rewrite_scene"

        rewrite_credit_cost = get_writer_credit_cost(rewrite_credit_action, rewrite_target_length)

        with col_left:
            st.subheader("📄 Editable Source Text")

            edited_source_text = st.text_area(
                "Edit selected source text before script rewrite:",
                value=selected_source_text,
                height=500,
                key=f"editable_rewrite_source_{scene_choice.replace(' ', '_')}_{st.session_state.manuscript_version}",
            )

            st.caption(f"Selected Source Words: {count_words(edited_source_text)}")
            st.markdown("---")

            credit_warning_ok = show_writer_credit_warning(rewrite_credit_action, rewrite_target_length)

            if st.button(
                f"Run Script Rewrite 🛠️ — Uses {rewrite_credit_cost} Credits",
                key="run_rewrite_engine",
                type="primary",
                use_container_width=True,
                disabled=not credit_warning_ok,
            ):
                require_writer_credits(rewrite_credit_action, rewrite_target_length)

                if not edited_source_text.strip():
                    st.error("No source text selected for script rewrite.")
                else:
                    with st.spinner("Running script rewrite..."):
                        try:
                            if st.session_state.get("demo_mode"):
                                user_row = refresh_demo_user_session()
                                ok, msg = is_demo_access_valid(user_row)
                                if not ok:
                                    st.error(msg)
                                    st.stop()

                            st.session_state.rewritten_manuscript = st.session_state.backend.run_rewrite_engine(
                                source_text=edited_source_text,
                                rewrite_mode=rewrite_mode,
                                scene_choice=scene_choice,
                                custom_instruction=custom_rewrite_instruction,
                                target_length=rewrite_target_length,
                                custom_target_word_count=custom_target_word_count,
                                tone_style=st.session_state.tone_style,
                                author_style=st.session_state.author_style,
                            )
                            st.session_state.rewrite_preview_id += 1

                            if ai_output_is_usable(st.session_state.rewritten_manuscript):
                                charge_writer_credits(
                                    rewrite_credit_action,
                                    st.session_state.rewritten_manuscript,
                                    rewrite_target_length,
                                )

                            if st.session_state.get("demo_mode"):
                                user_row = refresh_demo_user_session()
                                if user_row:
                                    tokens_used_now = max(1, len(str(st.session_state.rewritten_manuscript).split()))
                                    add_demo_token_usage(user_row["id"], tokens_used_now)

                                    user_row = refresh_demo_user_session()
                                    ok, msg = is_demo_access_valid(user_row)
                                    if not ok:
                                        st.warning(msg)
                                        st.stop()

                        except Exception as e:
                            st.session_state.rewritten_manuscript = (
                                "Script rewrite fallback output.\n\n"
                                f"Backend error: {e}\n\n"
                                "The script rewrite request was received, but the AI backend did not return usable text."
                            )
                            st.session_state.rewrite_preview_id += 1

                        st.rerun()

        st.markdown("---")
        st.subheader("✍️ Script Rewrite Output")

        st.text_area(
            "Script Rewrite Output:",
            value=st.session_state.rewritten_manuscript,
            height=450,
            label_visibility="collapsed",
            key=f"rewrite_output_preview_{st.session_state.rewrite_preview_id}",
        )

        preview_words = count_words(st.session_state.rewritten_manuscript)
        if st.session_state.rewritten_manuscript.strip():
            st.caption(f"Script Rewrite Preview Words: {preview_words}")

        if st.session_state.rewrite_applied_message:
            st.success(st.session_state.rewrite_applied_message)

        preview_is_safe_to_apply = (
            st.session_state.rewritten_manuscript.strip()
            and "Script rewrite fallback output" not in st.session_state.rewritten_manuscript
            and "Backend error:" not in st.session_state.rewritten_manuscript
        )

        if st.session_state.rewritten_manuscript.strip() and not preview_is_safe_to_apply:
            st.error("This script rewrite preview could not be used. Please run the script rewrite again.")

        if preview_is_safe_to_apply:
            if st.button(
                "Apply Script Rewrite to Full Manuscript ✅",
                key="apply_rewrite_to_manuscript",
                type="primary",
                use_container_width=True,
            ):
                cleaned_rewrite_for_manuscript = strip_internal_notices(st.session_state.rewritten_manuscript)

                if scene_choice == "Full Manuscript":
                    st.session_state.current_draft_output = cleaned_rewrite_for_manuscript
                    st.session_state.rewrite_applied_message = "Full manuscript replaced with script rewrite version."
                else:
                    beat_number = int(scene_choice.replace("Scene Beat ", ""))
                    st.session_state.current_draft_output = replace_scene_beat_text(
                        st.session_state.current_draft_output,
                        beat_number,
                        cleaned_rewrite_for_manuscript,
                    )
                    st.session_state.rewrite_applied_message = f"{scene_choice} script rewrite applied to the full manuscript."

                st.session_state.rewritten_manuscript = ""
                st.session_state.manuscript_version += 1
                save_manuscript_version_to_supabase(
                    st.session_state.current_draft_output,
                    version_label=f"Accepted Script Rewrite v{st.session_state.manuscript_version}",
                    source_stage="rewrite_engine",
                )
                st.session_state.rewrite_preview_id += 1

                save_result = save_current_project_session_to_supabase()
                if save_result:
                    st.session_state.rewrite_applied_message = (
                        st.session_state.rewrite_applied_message
                        + " Project saved to the current accepted state."
                    )
                else:
                    st.session_state.rewrite_applied_message = (
                        st.session_state.rewrite_applied_message
                        + " Project session was not saved automatically. Use Save Project before leaving."
                    )

                st.session_state.project_stage = "Isolation Guardrail"
                st.session_state.export_view = True
                st.rerun()

        st.markdown("---")

        try:
            final_doc = Document()
            final_style = final_doc.styles["Normal"]
            final_font = final_style.font
            final_font.name = "Times New Roman"
            final_font.size = Pt(12)

            final_export_text = strip_internal_notices(st.session_state.current_draft_output)

            for line in final_export_text.split("\n"):
                clean_line = line.strip()
                if clean_line:
                    if "CHAPTER SCENE BEAT" in clean_line:
                        p = final_doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(18)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run(clean_line)
                        run.bold = True
                    elif not clean_line.startswith("="):
                        p = final_doc.add_paragraph(clean_line)
                        p.paragraph_format.space_after = Pt(12)
                        p.paragraph_format.line_spacing = 1.15

            final_buffer = BytesIO()
            final_doc.save(final_buffer)
            final_docx_bytes = final_buffer.getvalue()

            st.download_button(
                label="Download Current Manuscript (.docx) 📥",
                data=final_docx_bytes,
                file_name="manuscript_current.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )

        except Exception as e:
            st.error(f"Failed to generate revised Word document: {e}")
